import pdfplumber
from docx import Document as DocxDocument
from pathlib import Path
from typing import List, Dict
import logging
import re

logger = logging.getLogger(__name__)


class TextExtractor:
    """Extract text from PDF and DOCX files."""

    @staticmethod
    async def extract(file_path: str, file_type: str) -> List[Dict]:
        """
        Extract text from a document.

        Returns:
            List of dicts with 'page_number' and 'text' keys.
        """
        if file_type == "pdf":
            return await TextExtractor._extract_pdf(file_path)
        elif file_type == "docx":
            return await TextExtractor._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    @staticmethod
    async def _extract_pdf(file_path: str) -> List[Dict]:
        """Extract text from PDF using pdfplumber."""
        pages = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    # Clean extracted text
                    text = TextExtractor._clean_text(text)
                    if text.strip():
                        pages.append({
                            "page_number": i + 1,
                            "text": text,
                        })
        except Exception as e:
            logger.error(f"Error extracting PDF {file_path}: {e}")
            raise

        return pages

    @staticmethod
    async def _extract_docx(file_path: str) -> List[Dict]:
        """Extract text from DOCX using python-docx."""
        try:
            doc = DocxDocument(file_path)
            full_text = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        full_text.append(" | ".join(row_text))

            combined_text = "\n\n".join(full_text)
            combined_text = TextExtractor._clean_text(combined_text)

            # DOCX doesn't have pages, so return as single "page"
            return [{"page_number": 1, "text": combined_text}]

        except Exception as e:
            logger.error(f"Error extracting DOCX {file_path}: {e}")
            raise

    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean and normalize extracted text."""
        # Remove excessive whitespace
        text = re.sub(r"\s+", " ", text)

        # Remove common header/footer patterns
        lines = text.split("\n")
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            # Skip lines that look like page numbers
            if re.match(r"^(\d+|page\s+\d+|- \d+ -)$", line, re.IGNORECASE):
                continue
            # Skip very short lines that might be headers/footers
            if len(line) < 5 and not any(c.isalpha() for c in line):
                continue
            cleaned_lines.append(line)

        text = "\n".join(cleaned_lines)

        # Normalize unicode
        text = text.encode("utf-8", errors="ignore").decode("utf-8")

        # Normalize line breaks
        text = re.sub(r"\n{3,}", "\n\n", text)

        return text.strip()
